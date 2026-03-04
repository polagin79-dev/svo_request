from aiogram import Router, F
from aiogram import Bot
from aiogram.filters import Command, CommandObject, StateFilter, and_f, or_f
from aiogram.types import Message, ReplyKeyboardRemove, FSInputFile
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State
from datetime import datetime, date
import re
import os

from keyboards.for_operator import zaiavki_period_kb
from handlers.stats_operator import StateZaiavki
from settings import Settings
from tools.safe_message import message_answer_safe, bot_send_photo_safe, bot_send_doc_safe

from mydb.db_work import getUnprocessedRequestText, getUnprocessedRequestPhoto, markProcessing, getPeriodText, getPeriodPhoto

import openpyxl
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Border, Side, PatternFill, Alignment

import docx
from docx import Document
from docx.shared import Inches

router = Router()

#text_operator = "/zaiavki - заявки\n" +\
#                "/zaiavki_period - заявки с какого-то числа\n" +\
#                "/zaiavki_photo - заявки в виде фото"

text_operator = "/zaiavki - заявки\n" +\
                "/zaiavki_period - заявки с какого-то числа\n"

@router.message(Command("operator"))  
async def cmd_operator(message: Message, bot: Bot, state: FSMContext):
        if message.chat.type == 'supergroup' or message.chat.type == 'group':
                return
        await message_answer_safe(
                message,
                text_operator,
                reply_markup=ReplyKeyboardRemove())
        await state.clear()

def removeUnnecessary(msg:str):
        return msg.replace('заявка','').replace('зявка','').replace('завка','').replace('заява','').replace('заяка','').replace('шт',' ').replace('шт.',' ').strip()

def dopSplit(text:str)->int:
        pattern=r"\b\d+\)"
        res = re.split(pattern, text)
        if res == None:
                res=[text]
        newres=[]
        for txt in res:
                if len(txt) == 0:
                        continue
                pattern=r'((\b[a-zA-Zа-яА-Я].+?)\s+(\d+))(\b|шт)'
                res2 = re.findall(pattern, txt)
                if res2 != None and len(res2)>0:
                        for row in res2:
                                newres.append([row[0],row[1],row[2]])
                else:
                        newres.append([txt,txt,1])
        return newres

def splitLines(msg:str)->list:
        ar = msg.split('\n')
        res=[]
        for row in ar:
                res=res+dopSplit(row)
        return res

def requestProcessingText(rows:list):
        # Открываем шаблонный файл
        wb = load_workbook(filename='TEMPLATE/template.xlsm', read_only=False, keep_vba=True)
        sheet = wb['заявки']
        # Записываем данные в ячейки
        i = 1
        for rec in rows:
                if rec['msg'] != None:
                        m=removeUnnecessary(rec['msg'])
                        msg_array = splitLines(m)
                        for msg in msg_array:
                                if len(msg) < 3:
                                        continue
                                utc_time = datetime.strptime(rec['dt_msg'], "%Y-%m-%d %H:%M:%S%z")
                                local_time = utc_time.astimezone()
                                sheet['A'+str(i+3)] = local_time.replace(tzinfo=None)
                                sheet['B'+str(i+3)] = rec['user_id']
                                sheet['C'+str(i+3)] = rec['full_name']
                                sheet['D'+str(i+3)] = rec['msg']
                                sheet['E'+str(i+3)] = msg[0]
                                if len(msg[1].strip())>3:
                                        sheet['G'+str(i+3)] = msg[1]
                                        sheet['H'+str(i+3)] = int(msg[2])
                                i=i+1
        sheet.freeze_panes = 'A4'
        # Границы
        thin_border = Border(left=Side(style='thin'),
                             right=Side(style='thin'),
                             top=Side(style='thin'),
                             bottom=Side(style='thin'))
        for r in range(4,i+3):
                for c in range(1,9):
                        sheet.cell(row=r, column=c).border = thin_border
        #обновляем сводные таблицы
        #pivot_sheet = wb["свод по товару"]
        #pivot = pivot_sheet._pivots[0]
        #pivot.cache.refreshOnLoad = True
        # Сохраняем файл на диск
        #wb.save('UPLOAD/requests.xlsm')
        # Открываем шаблонный файл
        #wb = load_workbook(filename='UPLOAD/requests.xlsm', read_only=False, keep_vba=True)
        #pivot_sheet = wb["свод по товару"]
        #for j in range(1,100):
        #        #print('requestProcessing: '+ 'A'+str(j), str(pivot_sheet['A'+str(j)].value))
        #        if len(str(pivot_sheet['A'+str(j)].value))<=3:
        #                continue
        #        if 'ozon.ru' in str(pivot_sheet['C'+str(j)].value):
        #                pivot_sheet['C'+str(j)]='озон'
        #        if 'wildberries.ru' in str(pivot_sheet['C'+str(j)].value):
        #                pivot_sheet['C'+str(j)]='валлберис'
                        
        #pivot_sheet = wb["свод по бойцу"]
        #pivot = pivot_sheet._pivots[0]
        #pivot.cache.refreshOnLoad = True
        
        # Сохраняем файл на диск
        wb.save('UPLOAD/requests.xlsm')        
        return [True, rows]

def inch_to_cm(value):
    return value / 2.54

def requestProcessingPhoto(rows:list):
        document = Document()
        #document = Document("UPLOAD/requests.docx")
        if rows:
                for rec in rows:
                        document.add_picture(rec['photo'], width=docx.shared.Inches(inch_to_cm(16)))
                        document.add_paragraph(rec['full_name'] + ': ' + rec['msg'])
                        document.add_paragraph(' ')
        else:
                document.add_paragraph('нет заявок')
        document.save("UPLOAD/requests.docx")
        return True

@router.message(Command("zaiavki"))  
async def cmd_zaiavki(message: Message, bot: Bot, command: CommandObject, state: FSMContext):
        if message.chat.type == 'supergroup' or message.chat.type == 'group':
                return
        #Читаем необработанные текстовые заявки из базы
        rows = getUnprocessedRequestText()
        if not rows:
                await message_answer_safe(
                        message,
                        "нет заявок",
                        reply_markup=ReplyKeyboardRemove()
                        )
                return
        if requestProcessingText(rows):
                if os.path.exists('UPLOAD/requests.xlsm'):                        
                        doc = FSInputFile('UPLOAD/requests.xlsm')
                        await bot_send_doc_safe(
                                Settings.bot,
                                message.from_user.id,
                                doc,
                                caption2="текстовые заявки",
                                reply_markup2 = ReplyKeyboardRemove()
                                )
                        #await bot.send_document(
                        #        message.from_user.id,
                        #        doc,
                        #        caption="тест",
                        #        reply_markup=ReplyKeyboardRemove()
                        #        )
                        markProcessing(rows, message.from_user.id, message.from_user.full_name)
                else:
                        await message_answer_safe(
                                message,
                                "неудалось сформировать заявки",
                                reply_markup=ReplyKeyboardRemove()
                                )

        #Читаем необработанные заявки в виде фото из базы
        rows = getUnprocessedRequestPhoto()
        if not rows:
                await message_answer_safe(
                        message,
                        "нет заявок",
                        reply_markup=ReplyKeyboardRemove()
                        )
                return
        if requestProcessingPhoto(rows):
                if os.path.exists('UPLOAD/requests.docx'):                        
                        doc = FSInputFile('UPLOAD/requests.docx')
                        await bot_send_doc_safe(
                                Settings.bot,
                                message.from_user.id,
                                doc,
                                caption2="заявки в виде фото",
                                reply_markup2 = ReplyKeyboardRemove()
                                )
                        #await bot.send_document(
                        #        message.from_user.id,
                        #        doc,
                        #        caption="тест",
                        #        reply_markup=ReplyKeyboardRemove()
                        #        )
                        markProcessing(rows, message.from_user.id, message.from_user.full_name)
                else:
                        await message_answer_safe(
                                message,
                                "неудалось сформировать заявки в виде фото",
                                reply_markup=ReplyKeyboardRemove()
                                )
        await state.clear()

@router.message(Command("zaiavki_photo"))  
async def cmd_zaiavki(message: Message, bot: Bot, command: CommandObject, state: FSMContext):
        if message.chat.type == 'supergroup' or message.chat.type == 'group':
                return
        #Читаем необработанные заявки из базы
        rows = getUnprocessedRequestPhoto()
        if not rows:
                await message_answer_safe(
                        message,
                        "нет заявок",
                        reply_markup=ReplyKeyboardRemove()
                        )
                return
        if requestProcessingPhoto(rows):
                if os.path.exists('UPLOAD/requests.docx'):                        
                        doc = FSInputFile('UPLOAD/requests.docx')
                        await bot_send_doc_safe(
                                Settings.bot,
                                message.from_user.id,
                                doc,
                                caption2="заявки",
                                reply_markup2 = ReplyKeyboardRemove()
                                )
                        #await bot.send_document(
                        #        message.from_user.id,
                        #        doc,
                        #        caption="тест",
                        #        reply_markup=ReplyKeyboardRemove()
                        #        )
                        markProcessing(rows, message.from_user.id, message.from_user.full_name)
                else:
                        await message_answer_safe(
                                message,
                                "неудалось сформировать заявки",
                                reply_markup=ReplyKeyboardRemove()
                                )
        await state.clear()

@router.message(Command("zaiavki_period"))  
async def cmd_load(message: Message, bot: Bot, command: CommandObject, state: FSMContext):
        if message.chat.type == 'supergroup' or message.chat.type == 'group':
                return
        await message_answer_safe(
                message,
                f"Укажите дату начиная с которой необходимо выгрузить заявки\n"+\
                "формат dd.mm.yy (например 02.11.25)",
                #parse_mode=ParseMode.HTML,
                reply_markup=zaiavki_period_kb()
                )
        await state.set_state(StateZaiavki.zaiavki_period)

@router.message(and_f(StateZaiavki.zaiavki_period, F.text == "ПРЕКРАТИТЬ"))
async def cmd_load(message: Message, bot: Bot, state: FSMContext):
        await message_answer_safe(
                message,
                text_operator,
                reply_markup=ReplyKeyboardRemove()
                )
        await state.clear()

@router.message(StateZaiavki.zaiavki_period)
async def cmd_load_period(message: Message, bot: Bot, state: FSMContext):
        ar=message.text.split('.')
        if len(ar) != 3:
                await message_answer_safe(
                        message,
                        "некорректная дата\nПовторите ввод",
                        reply_markup=zaiavki_period_kb()
                        )
                return
        if (not ar[0].isdigit()) or (not ar[1].isdigit()) or (not ar[0].isdigit()):
                await message_answer_safe(
                        message,
                        "некорректная дата\nПовторите ввод",
                        reply_markup=zaiavki_period_kb()
                        )
                return
        dd=int(ar[0])
        mm=int(ar[1])
        yy=int(ar[2])
        if (dd < 0) or (dd > 31) or (mm < 0) or (mm > 12) or (yy < 25) or ((yy > 35)and(yy<2025))or(yy>2035):
                await message_answer_safe(
                        message,
                        "некорректная дата\nПовторите ввод",
                        reply_markup=zaiavki_period_kb()
                        )
                return
        if yy < 100:
                yy=yy+2000
        d=date(yy,mm,dd)
        #Читаем заявки из базы с указанной даты
        rows=getPeriodText(d)
        if requestProcessingText(rows):
                if os.path.exists('UPLOAD/requests.xlsm'):                        
                        doc = FSInputFile('UPLOAD/requests.xlsm')
                        await bot.send_document(message.from_user.id, doc, caption="тест")
                else:
                        await message_answer_safe(message, "неудалось сформировать текстовые заявки")
        rows=getPeriodPhoto(d)
        if requestProcessingPhoto(rows):
                if os.path.exists('UPLOAD/requests.docx'):                        
                        doc = FSInputFile('UPLOAD/requests.docx')
                        await bot.send_document(message.from_user.id, doc, caption="тест")
                else:
                        await message_answer_safe(message, "неудалось сформировать заявки из фотографий")
        await state.clear()

def mainTest():
        requestProcessingText()
        #createXMLS()

if __name__ == "__main__":
    mainTest()
